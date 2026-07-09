import os
import re
from pathlib import Path
from tse.utils.helpers import strip_markdown
from tse.utils.logger import logger

def export_markdown(content: str, output_path: str) -> None:
    """Exports raw markdown content directly to a file."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    logger.info(f"Exported Markdown to {output_path}")

def export_text(content: str, output_path: str) -> None:
    """Strips markdown and exports plain text to a file."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    plain_text = strip_markdown(content)
    with open(path, "w", encoding="utf-8") as f:
        f.write(plain_text)
    logger.info(f"Exported plain text to {output_path}")

def export_word(content: str, output_path: str, title: str = "TSE Answer") -> None:
    """Exports content to a Word Document (.docx) with structured headings and lists."""
    try:
        import docx
    except ImportError:
        logger.error("python-docx is not installed. Word export failed.")
        raise ImportError("python-docx is required for Word export. Install with 'pip install python-docx'.")

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = docx.Document()
    doc.add_heading(title, 0)
    
    lines = content.splitlines()
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue
            
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r"^\d+\.\s+", stripped_line):
            doc.add_paragraph(stripped_line, style='List Number')
        else:
            doc.add_paragraph(stripped_line)
            
    doc.save(str(path))
    logger.info(f"Exported Word Document to {output_path}")

def export_pdf(content: str, output_path: str, title: str = "TSE Answer") -> None:
    """Exports content to a PDF document with standard layouts and styling."""
    try:
        from fpdf import FPDF
    except ImportError:
        logger.error("fpdf2 is not installed. PDF export failed.")
        raise ImportError("fpdf2 is required for PDF export. Install with 'pip install fpdf2'.")

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    class ExamPDF(FPDF):
        def header(self):
            # Header banner
            self.set_text_color(0, 168, 255) # Cyan hex representation
            self.set_font("Helvetica", "B", 10)
            self.cell(0, 10, "TSE CLI Academic Assistant", align="R")
            self.ln(10)
            self.set_draw_color(0, 168, 255)
            self.line(10, 18, 200, 18)
            
        def footer(self):
            # Footer banner
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    pdf = ExamPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    # Title Page/Header
    pdf.set_text_color(255, 0, 127) # Magenta
    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(0, 10, title, align="C")
    pdf.ln(10)
    
    # Body text settings
    pdf.set_text_color(30, 30, 30) # Dark Charcoal
    
    lines = content.splitlines()
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            pdf.ln(3)
            continue
            
        if line.startswith("# "):
            pdf.set_text_color(255, 0, 127) # Magenta
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 8, line[2:])
            pdf.ln(2)
        elif line.startswith("## "):
            pdf.set_text_color(0, 168, 255) # Cyan
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 6, line[3:])
            pdf.ln(2)
        elif line.startswith("### "):
            pdf.set_text_color(30, 30, 30)
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(0, 6, line[4:])
            pdf.ln(1)
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_text_color(30, 30, 30)
            pdf.set_font("Helvetica", "", 10)
            # Indent bullet point
            pdf.set_x(15)
            # Bullet symbol bullet character
            pdf.multi_cell(0, 6, f"* {line[2:]}")
        else:
            pdf.set_text_color(30, 30, 30)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 6, stripped_line)
            
    try:
        pdf.output(str(path))
        logger.info(f"Exported PDF Document to {output_path}")
    except Exception as e:
        logger.error(f"Failed to generate PDF: {e}")
        raise e
